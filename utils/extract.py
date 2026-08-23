"""Turns an uploaded document into plain text.

Supports PDF, DOCX, and plain-text/Markdown files. Anything else raises
UnsupportedFileError so the route can return a clean 400 to the caller.
"""
import io

import pdfplumber
from docx import Document

# Documents can be long; keep the amount we hand to the model bounded so a
# single upload can't blow the context window or the response time.
MAX_CHARS = 20000

TEXT_EXTENSIONS = {"txt", "md", "csv", "log"}


class UnsupportedFileError(Exception):
    pass


def file_extension(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def extract_text(filename: str, file_bytes: bytes) -> tuple[str, bool]:
    """Return (text, was_truncated) pulled from an uploaded document.

    Raises UnsupportedFileError for formats we don't handle, and ValueError
    if the file is empty or unreadable.
    """
    ext = file_extension(filename)

    if ext == "pdf":
        text = _extract_pdf(file_bytes)
    elif ext == "docx":
        text = _extract_docx(file_bytes)
    elif ext in TEXT_EXTENSIONS:
        text = _extract_plain_text(file_bytes)
    else:
        raise UnsupportedFileError(
            f"'.{ext}' isn't supported yet. Try a PDF, DOCX, TXT, or MD file."
        )

    text = text.strip()
    if not text:
        raise ValueError(
            "No readable text was found in that file. If it's a scanned "
            "PDF (images of pages rather than real text), text extraction "
            "won't pick anything up."
        )

    truncated = len(text) > MAX_CHARS
    return text[:MAX_CHARS], truncated


def _extract_pdf(file_bytes: bytes) -> str:
    pages = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text:
                pages.append(page_text)
    return "\n\n".join(pages)


def _extract_docx(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_plain_text(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Couldn't decode that file as text.")
